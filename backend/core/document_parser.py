"""
文档解析器

设计决策（全链路思考）：
1. 为什么支持多种格式？
   - 实际场景中用户可能上传 PDF、Word、Markdown、TXT 等各种格式
   - 每种格式有不同的解析挑战（PDF 有复杂布局，Word 有样式）

2. 为什么用 LangChain 的文档加载器？
   - 统一的 Document 接口（page_content + metadata）
   - 社区维护，兼容性好
   - 与 LangChain 生态无缝集成

3. 分块策略选择：
   - 评估结果显示 bge-m3 对不同 chunk_size 都有很好的检索效果
   - 默认使用 chunk_size=500, overlap=50（LangChain 推荐的起始值）
   - 后续可以根据评估结果调整
"""

import os
from pathlib import Path
from typing import Optional
from dataclasses import dataclass

from langchain_text_splitters import RecursiveCharacterTextSplitter
from backend.core.config import get_settings


@dataclass
class ParsedDocument:
    """解析后的文档"""
    doc_id: str
    filename: str
    file_type: str
    chunks: list[str]
    metadatas: list[dict]
    total_chars: int


class DocumentParser:
    def __init__(self):
        self.settings = get_settings()
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.settings.CHUNK_SIZE,
            chunk_overlap=self.settings.CHUNK_OVERLAP,
            length_function=len,
            separators=["\n\n", "\n", "。", ".", "！", "!", "？", "?", "；", ";", " "],
        )

    def parse(self, file_path: str, doc_id: str) -> Optional[ParsedDocument]:
        """解析文档文件"""
        path = Path(file_path)
        if not path.exists():
            return None

        suffix = path.suffix.lower()
        content = self._extract_text(file_path, suffix)
        if not content:
            return None

        # 分块
        chunks = self.text_splitter.split_text(content)

        # 构建元数据
        metadatas = [
            {
                "source": path.name,
                "chunk_index": i,
                "total_chunks": len(chunks),
            }
            for i in range(len(chunks))
        ]

        return ParsedDocument(
            doc_id=doc_id,
            filename=path.name,
            file_type=suffix,
            chunks=chunks,
            metadatas=metadatas,
            total_chars=len(content),
        )

    def _extract_text(self, file_path: str, suffix: str) -> Optional[str]:
        """根据文件类型提取文本"""
        try:
            if suffix == ".txt":
                return self._parse_txt(file_path)
            elif suffix == ".md":
                return self._parse_markdown(file_path)
            elif suffix == ".pdf":
                return self._parse_pdf(file_path)
            elif suffix == ".docx":
                return self._parse_docx(file_path)
            else:
                return None
        except Exception as e:
            print(f"Parse error ({suffix}): {e}")
            return None

    def _parse_txt(self, file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    def _parse_markdown(self, file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    def _parse_pdf(self, file_path: str) -> str:
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            texts = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    texts.append(f"[Page {i+1}]\n{text}")
            return "\n\n".join(texts)
        except ImportError:
            print("pypdf not installed, falling back to basic text extraction")
            return self._parse_txt(file_path)

    def _parse_docx(self, file_path: str) -> str:
        try:
            from docx import Document
            doc = Document(file_path)
            return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            print("python-docx not installed")
            return None


# 全局单例
_parser: Optional[DocumentParser] = None


def get_document_parser() -> DocumentParser:
    global _parser
    if _parser is None:
        _parser = DocumentParser()
    return _parser
